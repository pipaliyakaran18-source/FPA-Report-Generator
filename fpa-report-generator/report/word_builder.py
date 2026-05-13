"""Word (.docx) report builder.

Renders a five-page variance report:

    Page 1 – Cover page (company, period, analyst, CONFIDENTIAL banner)
    Page 2 – Variance table (cell shading, bold margin rows)
    Page 3 – Variance chart (Actual vs Budget + Variance % bars)
    Page 4 – Five-paragraph management commentary (bold lead sentences)
    Page 5 – Key ratios appendix

A footer on every page shows ``Company | CONFIDENTIAL | Page X``.
"""

from __future__ import annotations

from datetime import date
from io import BytesIO
from pathlib import Path

import matplotlib
matplotlib.use("Agg")  # headless backend — must come before pyplot import
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.document import Document as _DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.table import _Cell

from engine.calculator import FinancialAnalysis, LineItem
from engine.commentary import Commentary

GREEN_HEX = "C6EFCE"
RED_HEX = "F8CBAD"
NEUTRAL_HEX = "FFF2CC"
HEADER_HEX = "1F4E78"
LIGHT_GREY_HEX = "F2F2F2"


def _shade(cell: _Cell, hex_color: str) -> None:
    """Apply solid fill shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell: _Cell, color: str = "BFBFBF", size: str = "4") -> None:
    """Apply uniform thin borders to a single cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def _direction_color(direction: str) -> str:
    return {
        "Favourable": GREEN_HEX,
        "Adverse": RED_HEX,
        "Neutral": NEUTRAL_HEX,
    }[direction]


def _add_page_number_field(paragraph) -> None:
    """Insert a Word PAGE field that renders the current page number."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def _build_footer(doc: _DocumentType, company: str) -> None:
    """Add ``Company | CONFIDENTIAL | Page X`` to every page footer."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.text = ""

    # Tab stops: centre and right of the printable area (EMU is an integer).
    usable_width = int(section.page_width - section.left_margin - section.right_margin)
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(usable_width // 2, WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

    left_run = para.add_run(company)
    left_run.font.size = Pt(9)
    left_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    centre_run = para.add_run("\tCONFIDENTIAL")
    centre_run.bold = True
    centre_run.font.size = Pt(9)
    centre_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    right_run = para.add_run("\tPage ")
    right_run.font.size = Pt(9)
    right_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    _add_page_number_field(para)


def _set_margins(doc: _DocumentType) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def _add_cover_page(
    doc: _DocumentType,
    company: str,
    period: str,
    analyst: str,
) -> None:
    """Render the cover page and end with a page break."""
    # Top spacer
    for _ in range(4):
        doc.add_paragraph()

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = confidential.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(company)
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("FP&A Variance Report")
    sub_run.font.size = Pt(20)
    sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    period_p = doc.add_paragraph()
    period_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_run = period_p.add_run(period)
    period_run.font.size = Pt(18)
    period_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    period_run.bold = True

    for _ in range(8):
        doc.add_paragraph()

    analyst_p = doc.add_paragraph()
    analyst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    analyst_run = analyst_p.add_run(f"Prepared by  {analyst}")
    analyst_run.font.size = Pt(12)
    analyst_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(date.today().strftime("%d %B %Y"))
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _format_money(currency: str, value: float) -> str:
    sign = "-" if value < 0 else ""
    return f"{sign}{currency}{abs(value):,.1f}"


def _format_signed_money(currency: str, value: float) -> str:
    sign = "+" if value > 0 else "-" if value < 0 else ""
    return f"{sign}{currency}{abs(value):,.1f}"


def _add_section_heading(doc: _DocumentType, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    p.paragraph_format.space_after = Pt(6)


def _style_header_cell(cell: _Cell, text: str) -> None:
    cell.text = ""
    _shade(cell, HEADER_HEX)
    _set_cell_borders(cell, color="1F4E78")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _style_data_cell(
    cell: _Cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.RIGHT,
    fill: str | None = None,
    italic: bool = False,
) -> None:
    cell.text = ""
    _set_cell_borders(cell)
    if fill:
        _shade(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10)


def _add_variance_table(
    doc: _DocumentType,
    analysis: FinancialAnalysis,
    currency: str,
) -> None:
    """Variance table: 8 line items plus three bold margin-% rows."""
    line_items: list[LineItem] = analysis.line_items
    headers = [
        "Line Item",
        f"Actual ({currency}m)",
        f"Budget ({currency}m)",
        f"Variance ({currency}m)",
        "Variance %",
        "Direction",
    ]

    margin_rows = [
        (
            "Gross Margin %",
            analysis.gross_margin_actual,
            analysis.gross_margin_budget,
            analysis.gross_margin_delta_pp,
        ),
        (
            "EBITDA Margin %",
            analysis.ebitda_margin_actual,
            analysis.ebitda_margin_budget,
            analysis.ebitda_margin_delta_pp,
        ),
        (
            "Net Margin %",
            analysis.net_margin_actual,
            analysis.net_margin_budget,
            analysis.net_margin_delta_pp,
        ),
    ]

    total_rows = 1 + len(line_items) + len(margin_rows)
    table = doc.add_table(rows=total_rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for idx, header in enumerate(headers):
        _style_header_cell(table.rows[0].cells[idx], header)
    table.rows[0].height = Cm(0.8)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    row_idx = 1
    for item in line_items:
        row = table.rows[row_idx].cells
        zebra = LIGHT_GREY_HEX if row_idx % 2 == 0 else None
        _style_data_cell(row[0], item.name, align=WD_ALIGN_PARAGRAPH.LEFT, fill=zebra)
        _style_data_cell(row[1], f"{item.actual:,.1f}", fill=zebra)
        _style_data_cell(row[2], f"{item.budget:,.1f}", fill=zebra)
        var_fill = _direction_color(item.direction)
        _style_data_cell(
            row[3],
            _format_signed_money("", item.variance_amount),
            fill=var_fill,
            bold=True,
        )
        _style_data_cell(
            row[4],
            f"{item.variance_pct * 100:+.1f}%",
            fill=var_fill,
            bold=True,
        )
        _style_data_cell(
            row[5],
            item.direction,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=var_fill,
            bold=True,
        )
        row_idx += 1

    # Bold margin-% rows
    for label, actual, budget, delta_pp in margin_rows:
        row = table.rows[row_idx].cells
        if delta_pp > 0:
            fill = GREEN_HEX
        elif delta_pp < 0:
            fill = RED_HEX
        else:
            fill = NEUTRAL_HEX
        _style_data_cell(row[0], label, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, fill=fill)
        _style_data_cell(row[1], f"{actual * 100:.1f}%", bold=True, fill=fill)
        _style_data_cell(row[2], f"{budget * 100:.1f}%", bold=True, fill=fill)
        _style_data_cell(row[3], f"{delta_pp:+.1f} pp", bold=True, fill=fill)
        _style_data_cell(row[4], "—", bold=True, fill=fill, align=WD_ALIGN_PARAGRAPH.CENTER)
        direction = "Favourable" if delta_pp > 0 else "Adverse" if delta_pp < 0 else "Neutral"
        _style_data_cell(
            row[5],
            direction,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            fill=fill,
        )
        row_idx += 1


def _split_lead_sentence(text: str) -> tuple[str, str]:
    """Return ``(first_sentence_with_terminator, remainder)``."""
    text = text.strip()
    for terminator in (". ", "! ", "? "):
        idx = text.find(terminator)
        if idx != -1:
            return text[: idx + 1], text[idx + 2 :]
    return text, ""


def _add_commentary(doc: _DocumentType, commentary: Commentary) -> None:
    """Render the five paragraphs with bold first sentences."""
    for heading, body in commentary.as_paragraphs():
        h = doc.add_paragraph()
        h_run = h.add_run(heading)
        h_run.bold = True
        h_run.font.size = Pt(13)
        h_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)

        first, rest = _split_lead_sentence(body)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lead = p.add_run(first)
        lead.bold = True
        lead.font.size = Pt(11)
        if rest:
            tail = p.add_run(rest)
            tail.font.size = Pt(11)


def _render_variance_chart(
    analysis: FinancialAnalysis,
    currency: str,
    period: str,
) -> BytesIO:
    """Render an Actual vs Budget grouped bar chart to a PNG buffer.

    The lower panel shows the variance % per line item, coloured by
    favourable (green) or adverse (red) direction. Returns a seekable
    ``BytesIO`` ready for ``doc.add_picture``.
    """
    items = analysis.line_items
    labels = [li.name for li in items]
    actuals = np.array([li.actual for li in items])
    budgets = np.array([li.budget for li in items])
    variance_pct = np.array([li.variance_pct * 100 for li in items])
    bar_colours = [
        "#4CAF50" if li.direction == "Favourable"
        else "#E53935" if li.direction == "Adverse"
        else "#FBC02D"
        for li in items
    ]

    fig, (ax1, ax2) = plt.subplots(
        2, 1,
        figsize=(9.5, 6.2),
        gridspec_kw={"height_ratios": [3, 2], "hspace": 0.55},
    )

    # ---------- Top: Actual vs Budget grouped bars ----------
    x = np.arange(len(labels))
    width = 0.38
    ax1.bar(x - width / 2, actuals, width, label="Actual",
            color="#1F4E78", edgecolor="white")
    ax1.bar(x + width / 2, budgets, width, label="Budget",
            color="#BDD7EE", edgecolor="white")
    ax1.set_title(
        f"Actual vs Budget  —  {period}",
        fontsize=13, fontweight="bold", color="#1F4E78", pad=10,
    )
    ax1.set_ylabel(f"{currency} millions", fontsize=10)
    ax1.set_xticks(x)
    ax1.set_xticklabels(labels, rotation=20, ha="right", fontsize=9)
    ax1.legend(loc="upper right", frameon=False, fontsize=9)
    ax1.grid(axis="y", linestyle=":", alpha=0.4)
    ax1.spines["top"].set_visible(False)
    ax1.spines["right"].set_visible(False)
    for spine in ("left", "bottom"):
        ax1.spines[spine].set_color("#BFBFBF")

    # ---------- Bottom: Variance % bars ----------
    ax2.bar(x, variance_pct, color=bar_colours, edgecolor="white")
    ax2.axhline(0, color="#595959", linewidth=0.8)
    ax2.set_title(
        "Variance % vs Budget (colour = direction)",
        fontsize=11, fontweight="bold", color="#1F4E78", pad=8,
    )
    ax2.set_ylabel("Variance %", fontsize=10)
    ax2.set_xticks(x)
    ax2.set_xticklabels(labels, rotation=20, ha="right", fontsize=9)
    ax2.grid(axis="y", linestyle=":", alpha=0.4)
    ax2.spines["top"].set_visible(False)
    ax2.spines["right"].set_visible(False)
    for spine in ("left", "bottom"):
        ax2.spines[spine].set_color("#BFBFBF")
    for xi, v in zip(x, variance_pct):
        offset = 0.6 if v >= 0 else -1.4
        ax2.text(xi, v + offset, f"{v:+.1f}%", ha="center",
                 fontsize=8, color="#404040")

    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf


def _add_chart_page(
    doc: _DocumentType,
    analysis: FinancialAnalysis,
    currency: str,
    period: str,
) -> None:
    """Render the chart and place it on its own page."""
    _add_section_heading(doc, "Variance Visualisation")
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Top panel: actual versus budgeted values by line item. Bottom panel: "
        "variance percentage versus budget, shaded green where favourable and "
        "red where adverse."
    )
    intro_run.font.size = Pt(10)
    intro_run.italic = True
    intro_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    intro.paragraph_format.space_after = Pt(10)

    chart = _render_variance_chart(analysis, currency, period)
    pic_paragraph = doc.add_paragraph()
    pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_paragraph.add_run().add_picture(chart, width=Inches(6.5))

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _add_ratios_appendix(doc: _DocumentType, analysis: FinancialAnalysis) -> None:
    rows = [
        (
            "Gross Margin",
            analysis.gross_margin_actual,
            analysis.gross_margin_budget,
            analysis.gross_margin_delta_pp,
            True,
        ),
        (
            "EBITDA Margin",
            analysis.ebitda_margin_actual,
            analysis.ebitda_margin_budget,
            analysis.ebitda_margin_delta_pp,
            True,
        ),
        (
            "Net Margin",
            analysis.net_margin_actual,
            analysis.net_margin_budget,
            analysis.net_margin_delta_pp,
            True,
        ),
        (
            "Cost Ratio (COGS / Revenue)",
            analysis.cost_ratio_actual,
            analysis.cost_ratio_budget,
            analysis.cost_ratio_delta_pp,
            False,
        ),
    ]
    headers = ["Ratio", "Actual", "Budget", "Delta (pp)", "Direction"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, header in enumerate(headers):
        _style_header_cell(table.rows[0].cells[idx], header)

    for r_idx, (name, actual, budget, delta, higher_is_better) in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        if delta > 0:
            favourable = higher_is_better
        elif delta < 0:
            favourable = not higher_is_better
        else:
            favourable = None
        if favourable is True:
            fill, direction = GREEN_HEX, "Favourable"
        elif favourable is False:
            fill, direction = RED_HEX, "Adverse"
        else:
            fill, direction = NEUTRAL_HEX, "Neutral"
        _style_data_cell(cells[0], name, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
        _style_data_cell(cells[1], f"{actual * 100:.1f}%")
        _style_data_cell(cells[2], f"{budget * 100:.1f}%")
        _style_data_cell(cells[3], f"{delta:+.1f}", fill=fill, bold=True)
        _style_data_cell(
            cells[4],
            direction,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            fill=fill,
            bold=True,
        )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note_run = note.add_run(
        "Note: percentage-point (pp) deltas compare the period's actual ratio with the "
        "budgeted ratio. For the cost ratio, a lower outcome is treated as favourable."
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def build_report(
    *,
    analysis: FinancialAnalysis,
    commentary: Commentary,
    company: str,
    period: str,
    analyst: str,
    currency: str,
    output_path: Path | str,
) -> Path:
    """Render the four-page Word report and save to ``output_path``.

    Args:
        analysis: Output of :func:`engine.calculator.analyse`.
        commentary: Output of :func:`engine.commentary.generate`.
        company: Reporting entity name.
        period: Reporting period label (e.g. ``"Q1 FY2026"``).
        analyst: Author name shown on the cover page.
        currency: Currency symbol used in money columns.
        output_path: Where to save the ``.docx``. Parent directory is
            created if needed.

    Returns:
        The absolute path to the saved file.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_margins(doc)
    _build_footer(doc, company)

    # Default body font.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------- Page 1: Cover ----------
    _add_cover_page(doc, company, period, analyst)

    # ---------- Page 2: Variance Table ----------
    _add_section_heading(doc, f"Variance Analysis  |  {period}")
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        f"The table below sets out the actual results for {company} against "
        f"budget for {period}. Favourable variances are shaded green; "
        f"adverse variances are shaded red. All amounts are in millions "
        f"of {currency}."
    )
    intro_run.font.size = Pt(10)
    intro_run.italic = True
    intro_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    intro.paragraph_format.space_after = Pt(10)
    _add_variance_table(doc, analysis, currency)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------- Page 3: Chart ----------
    _add_chart_page(doc, analysis, currency, period)

    # ---------- Page 4: Commentary ----------
    _add_section_heading(doc, "Management Commentary")
    _add_commentary(doc, commentary)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------- Page 5: Ratios appendix ----------
    _add_section_heading(doc, "Appendix — Key Ratios")
    intro2 = doc.add_paragraph()
    intro2_run = intro2.add_run(
        "Summary of the profitability and cost ratios derived from the "
        "actual and budget figures for the period."
    )
    intro2_run.font.size = Pt(10)
    intro2_run.italic = True
    intro2_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    intro2.paragraph_format.space_after = Pt(10)
    _add_ratios_appendix(doc, analysis)

    doc.save(output_path)
    return output_path.resolve()
